"""Heading-aware document parsing — citations need structural breadcrumbs."""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

import markdown as md_lib
from bs4 import BeautifulSoup


# =============================================================================
# Data model
# =============================================================================
@dataclass
class Paragraph:
    text: str
    page: int
    section_path: list[str] = field(default_factory=list)  # ['H1', 'H2', 'H3']
    paragraph_index: int = 0


@dataclass
class ParsedDocument:
    name: str
    source_path: str
    paragraphs: list[Paragraph]
    page_count: int


# =============================================================================
# Markdown — preferred input format; heading structure is explicit
# =============================================================================
def parse_markdown(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8")
    html = md_lib.markdown(raw, extensions=["fenced_code", "tables"])
    soup = BeautifulSoup(html, "html.parser")

    paragraphs: list[Paragraph] = []
    section_stack: list[tuple[int, str]] = []  # (level, heading_text)
    para_idx = 0

    for el in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre"]):
        text = el.get_text(" ", strip=True)
        if not text:
            continue

        if el.name in {"h1", "h2", "h3", "h4"}:
            level = int(el.name[1])
            # Pop deeper or equal headings
            while section_stack and section_stack[-1][0] >= level:
                section_stack.pop()
            section_stack.append((level, text))
            continue

        section_path = [h for _, h in section_stack]
        paragraphs.append(
            Paragraph(
                text=text,
                page=1,
                section_path=section_path,
                paragraph_index=para_idx,
            )
        )
        para_idx += 1

    return ParsedDocument(
        name=path.name,
        source_path=str(path),
        paragraphs=paragraphs,
        page_count=1,
    )


# =============================================================================
# Plain text — heuristic heading detection (ALL CAPS or short lines, no period)
# =============================================================================
_HEADING_RE = re.compile(r"^[A-Z][A-Za-z0-9 &/\-]{2,80}$")


def _looks_like_heading(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 80:
        return False
    if line.endswith((".", ",", ";", ":", "?", "!")):
        return False
    # Either ALL CAPS or Title Case heading-y
    if line.isupper() and len(line) > 3:
        return True
    if _HEADING_RE.match(line) and line[0].isupper():
        words = line.split()
        # Title-case heuristic: most words start uppercase
        upper = sum(1 for w in words if w[:1].isupper())
        return upper / max(len(words), 1) >= 0.6
    return False


def parse_text(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8")
    paragraphs: list[Paragraph] = []
    section_stack: list[str] = []
    para_idx = 0

    # Split on blank lines
    blocks = re.split(r"\n\s*\n", raw)
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = block.splitlines()
        # Single-line block that looks like a heading
        if len(lines) == 1 and _looks_like_heading(lines[0]):
            heading = lines[0].strip()
            # Heuristic: ALL CAPS = top level, otherwise nest one level down
            if heading.isupper():
                section_stack = [heading]
            else:
                if section_stack:
                    section_stack = section_stack[:1] + [heading]
                else:
                    section_stack = [heading]
            continue

        # First line is heading, rest is body
        if len(lines) > 1 and _looks_like_heading(lines[0]):
            heading = lines[0].strip()
            if heading.isupper():
                section_stack = [heading]
            else:
                section_stack = section_stack[:1] + [heading] if section_stack else [heading]
            body = " ".join(l.strip() for l in lines[1:]).strip()
            if body:
                paragraphs.append(
                    Paragraph(text=body, page=1, section_path=list(section_stack), paragraph_index=para_idx)
                )
                para_idx += 1
            continue

        body = " ".join(l.strip() for l in lines).strip()
        if body:
            paragraphs.append(
                Paragraph(text=body, page=1, section_path=list(section_stack), paragraph_index=para_idx)
            )
            para_idx += 1

    return ParsedDocument(
        name=path.name,
        source_path=str(path),
        paragraphs=paragraphs,
        page_count=1,
    )


# =============================================================================
# PDF — page numbers come from the format; heading detection from font heuristics
# =============================================================================
def parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    paragraphs: list[Paragraph] = []
    section_stack: list[str] = []
    para_idx = 0

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        # Split into blocks (double newline) then check for headings
        blocks = re.split(r"\n\s*\n", text)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            lines = [l.strip() for l in block.splitlines() if l.strip()]
            if not lines:
                continue

            if len(lines) == 1 and _looks_like_heading(lines[0]):
                heading = lines[0]
                if heading.isupper():
                    section_stack = [heading]
                else:
                    section_stack = section_stack[:1] + [heading] if section_stack else [heading]
                continue

            if len(lines) > 1 and _looks_like_heading(lines[0]):
                heading = lines[0]
                if heading.isupper():
                    section_stack = [heading]
                else:
                    section_stack = section_stack[:1] + [heading] if section_stack else [heading]
                body = " ".join(lines[1:])
            else:
                body = " ".join(lines)

            if body:
                paragraphs.append(
                    Paragraph(text=body, page=page_num, section_path=list(section_stack), paragraph_index=para_idx)
                )
                para_idx += 1

    return ParsedDocument(
        name=path.name,
        source_path=str(path),
        paragraphs=paragraphs,
        page_count=len(reader.pages),
    )


# =============================================================================
# DOCX — uses style names directly for structure
# =============================================================================
def parse_docx(path: Path) -> ParsedDocument:
    from docx import Document

    doc = Document(str(path))
    paragraphs: list[Paragraph] = []
    section_stack: list[tuple[int, str]] = []
    para_idx = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
            while section_stack and section_stack[-1][0] >= level:
                section_stack.pop()
            section_stack.append((level, text))
            continue
        paragraphs.append(
            Paragraph(
                text=text,
                page=1,
                section_path=[h for _, h in section_stack],
                paragraph_index=para_idx,
            )
        )
        para_idx += 1

    return ParsedDocument(
        name=path.name,
        source_path=str(path),
        paragraphs=paragraphs,
        page_count=1,
    )


# =============================================================================
# Dispatch
# =============================================================================
PARSERS = {
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_text,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
}


def parse_any(path: Path) -> ParsedDocument:
    ext = path.suffix.lower()
    if ext not in PARSERS:
        raise ValueError(f"Unsupported file type: {ext}")
    return PARSERS[ext](path)


def iter_sources(source_dir: Path) -> Iterator[Path]:
    for p in sorted(source_dir.rglob("*")):
        if p.is_file() and p.suffix.lower() in PARSERS:
            yield p
